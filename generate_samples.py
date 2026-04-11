import os
import docx

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLES_DIR = os.path.join(BASE_DIR, "frontend", "samples")

os.makedirs(SAMPLES_DIR, exist_ok=True)

samples = {
    "Sample_A_Original.docx": "Artificial intelligence (AI) has rapidly transformed modern society. From healthcare diagnostics to autonomous vehicles, machine learning models are being deployed to solve complex problems. However, this widespread adoption brings challenges regarding ethics, bias, and data privacy. Researchers argue that without proper regulatory frameworks, AI systems might exacerbate existing inequalities. Therefore, developing explainable AI methods and ensuring transparency in model training is crucial for the future of algorithmic decision-making.",
    
    "Sample_B_CopyPaste.docx": "Artificial intelligence (AI) has rapidly transformed modern society. From healthcare diagnostics to autonomous vehicles, machine learning models are being deployed to solve complex problems. However, this widespread adoption brings challenges regarding ethics, bias, and data privacy. Researchers argue that without proper regulatory frameworks, AI systems might exacerbate existing inequalities. Therefore, developing explainable AI methods and ensuring transparency in model training is crucial for the future of algorithmic decision-making.",
    
    "Sample_C_Paraphrased.docx": "The modern world has been significantly changed by artificial intelligence. Machine learning algorithms are now utilized to tackle complicated issues in fields ranging from self-driving cars to medical evaluations. Yet, this extensive use introduces significant concerns about privacy, fairness, and morality. Experts point out that without adequate governance, automated systems could widen current social divides. Thus, creating interpretable algorithms and maintaining openness during the development process is essential for tomorrow's automated choices.",
    
    "Sample_D_Different.docx": "Photosynthesis is the process used by plants, algae and certain bacteria to harness energy from sunlight and turn it into chemical energy. Here, the energy from light is used to convert water, carbon dioxide, and minerals into oxygen and energy-rich organic compounds. This process is fundamental to life on Earth as it maintains the oxygen levels in the atmosphere and supplies most of the biological energy necessary for complex life."
}

for filename, text in samples.items():
    doc = docx.Document()
    doc.add_heading(filename.replace(".docx", ""), 0)
    doc.add_paragraph(text)
    file_path = os.path.join(SAMPLES_DIR, filename)
    doc.save(file_path)
    print(f"Generated {file_path}")
